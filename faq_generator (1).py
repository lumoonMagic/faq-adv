import streamlit as st
from docx import Document
from docx.shared import Inches
import tempfile
import json
import os
import google.generativeai as genai
from supabase import create_client

# --- CONFIG ---
SUPABASE_URL = st.secrets["SUPABASE_URL"]
SUPABASE_KEY = st.secrets["SUPABASE_KEY"]
GEMINI_API_KEY = st.secrets["GEMINI_API_KEY"]

supabase = create_client(SUPABASE_URL, SUPABASE_KEY)
genai.configure(api_key=GEMINI_API_KEY)

# --- FUNCTIONS ---
def load_faqs():
    response = supabase.table("faqs").select("*").limit(1).execute()
    if response.data:
        row = response.data[0]
        return row["data"]["faqs"], row["id"]
    else:
        insert_res = supabase.table("faqs").insert({"data": {"faqs": []}}).execute()
        new_id = insert_res.data[0]["id"]
        return [], new_id

def save_faqs(faqs, row_id):
    supabase.table("faqs").update({"data": {"faqs": faqs}}).eq("id", row_id).execute()

def validate_steps_with_gemini(question, steps):
    prompt = f"""
You are an expert technical documentation assistant. Review the following steps for the FAQ question: "{question}". 

1️⃣ Highlight if the question is addressed in the steps.  
2️⃣ Suggest alternatives or missing steps for clarity.  
3️⃣ Return a cleaned and improved version of the steps.

Steps:
{steps}
"""
    model = genai.GenerativeModel("gemini-2.5-flash")
    response = model.generate_content(prompt)
    return response.text

# --- LOAD DATA ---
if 'faq_list' not in st.session_state or 'faq_row_id' not in st.session_state:
    faqs, row_id = load_faqs()
    st.session_state['faq_list'] = faqs
    st.session_state['faq_row_id'] = row_id

# --- UI ---
st.title("📄 Troubleshooting — FAQ Generator")

assignees = list(set([faq["assignee"] for faq in st.session_state['faq_list']])) or ["No Assignee"]
selected_assignee = st.selectbox("👤 Select Assignee", assignees)

filtered_faqs = [faq for faq in st.session_state['faq_list'] if faq["assignee"] == selected_assignee]

faq_options = [faq["question"] for faq in filtered_faqs] or ["No FAQ"]
selected_faq = st.selectbox("❓ Select FAQ", faq_options)

# --- ADD FAQ ---
with st.form("add_faq_form"):
    new_q = st.text_input("➕ New FAQ Question")
    new_a = st.text_input("Assign to")
    submitted = st.form_submit_button("Add FAQ")
    if submitted:
        if new_q and new_a:
            new_faq = {"question": new_q, "assignee": new_a}
            st.session_state['faq_list'].append(new_faq)
            save_faqs(st.session_state['faq_list'], st.session_state['faq_row_id'])
            st.success(f"Added: {new_q}")
        else:
            st.warning("Please provide both a question and assignee.")

# --- DELETE FAQ ---
if selected_faq != "No FAQ":
    if st.button(f"🗑️ Delete '{selected_faq}'"):
        updated_list = [faq for faq in st.session_state['faq_list'] if faq["question"] != selected_faq]
        st.session_state['faq_list'] = updated_list
        save_faqs(updated_list, st.session_state['faq_row_id'])
        st.success(f"Deleted: {selected_faq}")

# --- BUILD FAQ CONTENT ---
st.subheader(f"📌 Generating FAQ for: {selected_faq}")
summary = st.text_area("Summary")

# Steps
if 'steps' not in st.session_state:
    st.session_state['steps'] = []

if st.button("Add Step"):
    st.session_state['steps'].append({"text": "", "screenshot": None, "query": ""})

for i, step in enumerate(st.session_state['steps']):
    st.session_state['steps'][i]["text"] = st.text_input(f"Step {i+1} Text", value=step["text"], key=f"step_text_{i}")
    st.session_state['steps'][i]["screenshot"] = st.file_uploader(f"Step {i+1} Screenshot", type=['png', 'jpg', 'jpeg'], key=f"step_ss_{i}")
    st.session_state['steps'][i]["query"] = st.text_area(f"Step {i+1} Query Template", value=step["query"], key=f"step_q_{i}")

notes = st.text_area("Additional Notes")

# --- GENERATE DOC ---
if st.button("📄 Generate FAQ Document"):
    doc = Document()
    doc.add_heading('FAQ Document', level=1)
    doc.add_heading('Question', level=2)
    doc.add_paragraph(selected_faq)
    doc.add_heading('Summary', level=2)
    doc.add_paragraph(summary)
    doc.add_heading('Steps', level=2)

    for idx, step in enumerate(st.session_state['steps']):
        doc.add_paragraph(f"Step {idx+1}: {step['text']}")
        if step["query"]:
            doc.add_paragraph(f"Query Template: {step['query']}")
        if step["screenshot"]:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmpfile:
                tmpfile.write(step["screenshot"].read())
                tmpfile.flush()
                doc.add_picture(tmpfile.name, width=Inches(4))

    doc.add_heading('Additional Notes', level=2)
    doc.add_paragraph(notes)

    tmp_out = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(tmp_out.name)
    st.download_button("Download FAQ Document", data=open(tmp_out.name, 'rb').read(),
                       file_name='FAQ_Generated.docx',
                       mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

# --- GEMINI VALIDATION ---
if st.button("🤖 Validate Steps (Gemini)"):
    steps_text = "\n".join([f"Step {idx+1}: {s['text']}" for idx, s in enumerate(st.session_state['steps'])])
    with st.spinner("Validating with Gemini..."):
        validation = validate_steps_with_gemini(selected_faq, steps_text)
    st.subheader("Gemini Validation Result")
    st.write(validation)
    tmp_out_val = tempfile.NamedTemporaryFile(delete=False, suffix='.txt')
    tmp_out_val.write(validation.encode())
    tmp_out_val.flush()
    st.download_button("Download Validation Report", data=open(tmp_out_val.name, 'rb').read(),
                       file_name='FAQ_Validation.txt',
                       mime='text/plain')
