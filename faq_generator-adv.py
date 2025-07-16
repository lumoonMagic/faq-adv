import streamlit as st
from docx import Document
from docx.shared import Inches
import tempfile
import zipfile
import io
import os
import google.generativeai as genai
from supabase import create_client

# --- CONFIG ---
SUPABASE_URL = st.secrets["SUPABASE_URL"]
SUPABASE_KEY = st.secrets["SUPABASE_KEY"]
GEMINI_API_KEY = st.secrets["GEMINI_API_KEY"]

supabase = create_client(SUPABASE_URL, SUPABASE_KEY)
genai.configure(api_key=GEMINI_API_KEY)

# --- FUNCTIONS ---
def load_faqs():
    response = supabase.table("faqs_adv").select("*").execute()
    return response.data if response.data else []

def validate_steps_with_gemini(question, steps):
    prompt = f"""
You are an expert technical documentation assistant. Review the following steps for the FAQ question: "{question}". 

1. Highlight if the question is addressed in the steps.  
2. Suggest alternatives or missing steps for clarity.  
3. Return a cleaned and improved version of the steps.

Steps:
{steps}
"""
    model = genai.GenerativeModel("gemini-2.5-flash")
    response = model.generate_content(prompt)
    return response.text

# --- LOAD DATA ---
if "faq_data" not in st.session_state:
    st.session_state.faq_data = load_faqs()

# --- UI ---
st.title("📄 Troubleshooting — FAQ Generator")

assignees = sorted(list(set([faq["data"].get("assignee", "") for faq in st.session_state.faq_data if "data" in faq])))
selected_assignee = st.selectbox("👤 Select Assignee", assignees)

faq_options = [faq for faq in st.session_state.faq_data if faq["data"].get("assignee") == selected_assignee]
faq_map = {faq["data"]["question"]: faq for faq in faq_options}
questions = list(faq_map.keys())
selected_q = st.selectbox("❓ Select FAQ", questions)
faq_entry = faq_map.get(selected_q)
faq_data = faq_entry.get("data", {}) if faq_entry else {}

# --- Inputs ---
st.subheader(f"📌 Editing FAQ: {selected_q}")
summary = st.text_area("[Summary]", value=faq_data.get("content", {}).get("summary", ""))
notes = st.text_area("[Additional Notes]", value=faq_data.get("content", {}).get("notes", ""))
keywords_input = st.text_input("Keywords (comma-separated)", value=", ".join(faq_entry.get("keywords", [])))
keywords_list = [k.strip() for k in keywords_input.split(",") if k.strip()]

if "steps" not in st.session_state:
    st.session_state["steps"] = faq_data.get("content", {}).get("steps", [])

if st.button("➕ Add Step"):
    st.session_state["steps"].append({"text": "", "screenshot": None, "query": ""})

st.markdown("---")
st.subheader("🪜 Steps")
if "pending_screenshots" not in st.session_state:
    st.session_state["pending_screenshots"] = {}

for i, step in enumerate(st.session_state["steps"]):
    st.session_state["steps"][i]["text"] = st.text_input(f"Step {i+1} Text", value=step.get("text", ""), key=f"step_text_{i}")
    st.session_state["steps"][i]["query"] = st.text_area(f"Step {i+1} Query", value=step.get("query", ""), key=f"step_query_{i}")
    uploaded = st.file_uploader(f"Step {i+1} Screenshot", type=["png", "jpg", "jpeg"], key=f"ss_{i}")
    if uploaded:
        st.session_state["pending_screenshots"][i+1] = uploaded

# --- Generate DOC + ZIP ---
if st.button("📄 Generate FAQ Document"):
    doc = Document()
    doc.add_heading("FAQ Document", level=1)
    doc.add_paragraph("[Question]")
    doc.add_paragraph(selected_q)
    doc.add_paragraph("[Summary]")
    doc.add_paragraph(summary)
    doc.add_paragraph("[Steps]")

    zip_buffer = io.BytesIO()
    screenshot_found = False
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for idx, step in enumerate(st.session_state["steps"]):
            doc.add_paragraph(f"[Step {idx+1}]")
            doc.add_paragraph(step.get("text", ""))
            if step.get("query"):
                doc.add_paragraph("[Query Template]")
                doc.add_paragraph(step.get("query"))
            if step.get("screenshot"):
                screenshot_found = True
                with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmpfile:
                    tmpfile.write(step["screenshot"].read())
                    tmpfile.flush()
                    doc.add_paragraph("[Screenshot]")
                    doc.add_picture(tmpfile.name, width=Inches(4))
                    zip_file.write(tmpfile.name, arcname=f"Step{idx+1}_screenshot.png")

    doc.add_paragraph("[Additional Notes]")
    doc.add_paragraph(notes)
    tmp_out = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp_out.name)

    with open(tmp_out.name, 'rb') as doc_file:
        st.download_button("Download FAQ Document", data=doc_file.read(), file_name="FAQ_Generated.docx")
    if screenshot_found:
        zip_buffer.seek(0)
        st.download_button("Download Screenshots ZIP", data=zip_buffer.read(), file_name="FAQ_Screenshots.zip", mime="application/zip")

# --- Validate Steps ---
if st.button("🧠 Validate Steps with Gemini"):
    step_text = "\n".join([f"Step {idx+1}: {s['text']}" for idx, s in enumerate(st.session_state["steps"])])
    with st.spinner("Validating..."):
        result = validate_steps_with_gemini(selected_q, step_text)
    st.code(result)

# --- Save to DB ---
if st.button("💾 Save / Update FAQ in DB"):
    for step_num, file in st.session_state['pending_screenshots'].items():
        st.session_state['steps'][step_num-1]["screenshot"] = file

    updated_data = {
        "question": selected_q,
        "assignee": faq_entry["data"].get("assignee"),
        "content": {
            "summary": summary,
            "steps": st.session_state["steps"],
            "notes": notes
        }
    }

    supabase.table("faqs_adv").update({
        "data": updated_data,
        "title": selected_q,
        "question": selected_q,
        "keywords": keywords_list
    }).eq("id", faq_entry["id"]).execute()

    st.success("✅ FAQ updated successfully!")
